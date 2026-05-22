"""
Document parsers — convert PDF, DOCX, XLSX, CSV, TXT, Markdown to plain text.
"""

import csv
import io
from pathlib import Path
from typing import BinaryIO

import structlog

log = structlog.get_logger()

SUPPORTED_TYPES = {"pdf", "docx", "xlsx", "csv", "txt", "md", "markdown"}


class UnsupportedFileTypeError(ValueError):
    pass


def detect_file_type(filename: str) -> str:
    suffix = Path(filename).suffix.lower().lstrip(".")
    if suffix == "markdown":
        return "md"
    if suffix not in SUPPORTED_TYPES:
        raise UnsupportedFileTypeError(f"Unsupported file type: {suffix}")
    return suffix


def parse(file: BinaryIO, file_type: str) -> str:
    """Parse a binary file into plain text. Raises UnsupportedFileTypeError."""
    file_type = file_type.lower()
    if file_type == "pdf":
        return _parse_pdf(file)
    if file_type == "docx":
        return _parse_docx(file)
    if file_type == "xlsx":
        return _parse_xlsx(file)
    if file_type == "csv":
        return _parse_csv(file)
    if file_type in ("txt", "md", "markdown"):
        return _parse_text(file)
    raise UnsupportedFileTypeError(f"No parser for {file_type}")


def _parse_pdf(file: BinaryIO) -> str:
    from pypdf import PdfReader

    reader = PdfReader(file)
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as e:
            log.warning("pdf_page_extract_failed", page=i, error=str(e))
            text = ""
        if text.strip():
            pages.append(text.strip())
    return "\n\n".join(pages)


def _parse_docx(file: BinaryIO) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(file)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Include tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)
    return "\n\n".join(paragraphs)


def _parse_xlsx(file: BinaryIO) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(file, read_only=True, data_only=True)
    sections = []
    for sheet in wb.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cell.strip() for cell in cells):
                rows.append(" | ".join(cells))
        if rows:
            sections.append(f"# Sheet: {sheet.title}\n\n" + "\n".join(rows))
    return "\n\n".join(sections)


def _parse_csv(file: BinaryIO) -> str:
    raw = file.read().decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(raw))
    lines = []
    for row in reader:
        if any(cell.strip() for cell in row):
            lines.append(" | ".join(row))
    return "\n".join(lines)


def _parse_text(file: BinaryIO) -> str:
    raw = file.read()
    if isinstance(raw, bytes):
        return raw.decode("utf-8", errors="replace")
    return raw
